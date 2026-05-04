"""Парсинг Word (.docx)."""

from __future__ import annotations

import asyncio
from decimal import Decimal
from pathlib import Path

from docx import Document

from core.llm.client import LLMClient, LLMParseError
from core.parsers.types import ParseResult, ParsedItem


class DocxParser:
    def _extract_text_sync(self, file_path: Path) -> str:
        doc = Document(file_path)
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    async def parse_with_claude(self, text: str) -> ParseResult:
        client = LLMClient()
        prompt = f"""Извлеки ВСЕ позиции из документа в JSON.
Если количество в виде «27.34 м» / «3 изделия» — извлекай и число, и единицу.
Если qty не указано — поставь 1.

Формат: {{"items": [{{"original_text": "...", "suggested_name": "...", "suggested_description": "...", "quantity": число, "unit": "шт."}}]}}
Только JSON, без markdown.

ТЕКСТ:
{text[:120_000]}"""
        try:
            data = await client.complete_json(
                "Ты извлекаешь строки спецификации из Word.",
                prompt,
                max_tokens=16000,
                temperature=0,
            )
        except (LLMParseError, Exception) as e:  # noqa: BLE001
            return ParseResult(items=[], confidence=0.0, needs_manual_review=True, parser_notes=[str(e)])
        out: list[ParsedItem] = []
        for row in data.get("items") or []:
            q = row.get("quantity")
            qty = Decimal(str(q)) if q is not None else None
            out.append(
                ParsedItem(
                    original_text=str(row.get("original_text", "")),
                    suggested_name=str(row.get("suggested_name", "")),
                    suggested_description=str(row.get("suggested_description", "")),
                    quantity=qty,
                    unit=str(row.get("unit") or "шт."),
                    raw_data={},
                )
            )
        return ParseResult(items=out, confidence=0.65, needs_manual_review=len(out) == 0, parser_notes=[])

    async def parse(self, file_path: Path) -> ParseResult:
        text = await asyncio.to_thread(self._extract_text_sync, file_path)
        if not text.strip():
            return ParseResult(items=[], confidence=0.0, needs_manual_review=True, parser_notes=["Пустой docx"])
        return await self.parse_with_claude(text)
